"""
文件处理器组件
基于marker组件实现的文件到markdown转换器
支持多种文件格式，提供高质量的markdown输出
"""

import asyncio
import os
import time
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from fastapi import HTTPException, UploadFile
from loguru import logger

from examples.conf.file_extractor_config import (
    ConfigManager,
    MarkdownContent,
    MarkdownExtractor,
    MarkerConfig,
)
from examples.image_analyzer import ImageAnalyzer, default_analyzer


class FileProcessor:
    """文件处理器主类"""

    def __init__(
        self,
        config: Optional[MarkerConfig] = None,
        upload_dir: str = "uploads",
        image_analyzer: Optional[ImageAnalyzer] = None,
    ):
        self.config_manager = ConfigManager(config)
        self.markdown_extractor = MarkdownExtractor()
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)

        # 图片分析器
        self.image_analyzer = image_analyzer or default_analyzer

        # 初始化marker组件（延迟导入以避免依赖问题）
        self._converter = None
        self._model_dict = None

    def _init_marker_components(self):
        """初始化marker组件"""
        if self._converter is None:
            try:
                from marker.config.parser import ConfigParser
                from marker.converters.pdf import PdfConverter
                from marker.models import create_model_dict

                # 创建配置解析器
                config_dict = self.config_manager.get_config_dict()
                config_parser = ConfigParser(config_dict)

                # 创建模型字典
                self._model_dict = create_model_dict()

                # 创建转换器
                self._converter = PdfConverter(
                    config=config_dict,
                    artifact_dict=self._model_dict,
                    processor_list=config_parser.get_processors(),
                    renderer=config_parser.get_renderer(),
                )

                # 如果使用LLM，设置LLM服务
                if config_dict.get("use_llm", False):
                    llm_service = config_parser.get_llm_service()
                    if llm_service:
                        self._converter.llm_service = llm_service

                self._marker_available = True

            except ImportError as e:
                print(f"⚠️ Marker依赖不可用，将使用基础文本提取: {str(e)}")
                self._marker_available = False
                self._converter = None
            except Exception as e:
                print(f"⚠️ Marker初始化失败，将使用基础文本提取: {str(e)}")
                self._marker_available = False
                self._converter = None

    async def process_file(self, file_path: Union[str, Path]) -> MarkdownContent:
        """处理单个文件"""
        start_time = time.time()
        file_path = Path(file_path)

        logger.info(f"📁 开始处理文件: {file_path.name}")
        logger.debug(f"文件路径: {file_path}")

        # 检查文件是否存在
        if not file_path.exists():
            logger.error(f"❌ 文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")

        logger.debug(f"✅ 文件存在检查通过")

        # 检查文件类型支持
        if not self.config_manager.is_file_supported(file_path.name):
            logger.error(f"❌ 不支持的文件类型: {file_path.suffix}")
            raise ValueError(f"不支持的文件类型: {file_path.suffix}")

        logger.debug(f"✅ 文件类型检查通过: {file_path.suffix}")

        # 验证文件大小
        file_size = file_path.stat().st_size
        file_size_mb = file_size / 1024 / 1024
        max_size_mb = self.config_manager.config.max_file_size / 1024 / 1024

        logger.info(f"📊 文件大小: {file_size_mb:.2f}MB (限制: {max_size_mb:.1f}MB)")

        if file_size > self.config_manager.config.max_file_size:
            logger.error(
                f"❌ 文件大小超过限制: {file_size_mb:.2f}MB > {max_size_mb:.1f}MB"
            )
            raise ValueError(f"文件大小超过限制 ({max_size_mb:.1f}MB)")

        logger.debug(f"✅ 文件大小检查通过")

        # 检查文件扩展名，如果是Markdown文件，直接处理
        file_extension = file_path.suffix.lower()
        if file_extension == ".md":
            logger.info(f"🔍 检测到Markdown文件，使用直接处理模式")
            try:
                markdown_content = await self._process_markdown_directly(file_path)
                processing_time = time.time() - start_time
                logger.success(
                    f"✅ Markdown文件处理完成，耗时: {processing_time:.2f}秒"
                )
                return markdown_content
            except Exception as e:
                logger.error(f"❌ Markdown文件直接处理失败: {str(e)}")
                raise Exception(f"Markdown文件处理失败: {str(e)}")

        # 对于非Markdown文件，使用原有的处理流程
        logger.info(f"🔧 使用标准处理流程处理文件: {file_extension}")

        try:
            # 初始化marker组件
            logger.debug("🔄 初始化Marker组件...")
            self._init_marker_components()

            # 检查marker是否可用
            if (
                hasattr(self, "_marker_available")
                and self._marker_available
                and self._converter
            ):
                logger.info("🚀 使用Marker进行文件转换")
                # 使用marker转换文件
                rendered_result = await self._convert_with_marker(str(file_path))
                # 提取markdown内容
                markdown_content = self.markdown_extractor.extract_from_rendered(
                    rendered_result
                )
                logger.success("✅ Marker转换完成")

                # 如果有图片，进行图片分析
                if markdown_content.images and self.image_analyzer.is_available():
                    logger.info("🖼️ 检测到图片，开始进行图片分析...")
                    markdown_content = await self._analyze_images_in_content(
                        markdown_content
                    )

            else:
                logger.warning("⚠️ Marker不可用，使用基础文本提取")
                # 使用基础文本提取作为后备方案
                markdown_content = await self._fallback_text_extraction(file_path)

            processing_time = time.time() - start_time
            logger.success(f"✅ 文件处理完成，耗时: {processing_time:.2f}秒")
            return markdown_content

        except Exception as e:
            logger.error(f"❌ Marker处理失败: {str(e)}")
            # 如果marker失败，尝试基础提取
            try:
                logger.info("🔄 尝试基础文本提取作为后备方案")
                markdown_content = await self._fallback_text_extraction(file_path)
                processing_time = time.time() - start_time
                logger.success(f"✅ 基础文本提取完成，耗时: {processing_time:.2f}秒")
                return markdown_content
            except Exception as fallback_error:
                processing_time = time.time() - start_time
                logger.error(f"❌ 所有处理方法都失败，总耗时: {processing_time:.2f}秒")
                raise Exception(
                    f"文件处理失败 - Marker: {str(e)}, 基础提取: {str(fallback_error)}"
                )

    async def _convert_with_marker(self, file_path: str):
        """使用marker转换文件"""
        logger.debug(f"🔄 开始Marker转换: {file_path}")
        start_time = time.time()

        try:
            # 在线程池中运行marker转换（因为marker是同步的）
            logger.debug("⚡ 在线程池中执行Marker转换...")
            loop = asyncio.get_event_loop()
            rendered_result = await loop.run_in_executor(
                None, self._converter, file_path
            )

            conversion_time = time.time() - start_time
            logger.debug(f"✅ Marker转换完成，耗时: {conversion_time:.2f}秒")
            return rendered_result

        except Exception as e:
            conversion_time = time.time() - start_time
            logger.error(
                f"❌ Marker转换失败，耗时: {conversion_time:.2f}秒，错误: {str(e)}"
            )
            raise Exception(f"Marker转换失败: {str(e)}")

    async def _fallback_text_extraction(self, file_path: Path):
        """基础文本提取作为后备方案"""
        logger.info(f"🔧 开始基础文本提取: {file_path.name}")
        start_time = time.time()

        file_extension = file_path.suffix.lower()
        content = ""

        logger.debug(f"📋 文件类型: {file_extension}")

        try:
            if file_extension == ".pdf":
                logger.debug("📄 使用PDF提取器...")
                content = await self._extract_pdf_basic(file_path)
            elif file_extension in [".docx", ".doc"]:
                logger.debug("📝 使用Word文档提取器...")
                content = await self._extract_docx_basic(file_path)
            elif file_extension == ".txt":
                logger.debug("📃 使用文本文件提取器...")
                content = await self._extract_txt_basic(file_path)
            elif file_extension == ".md":
                logger.debug("📝 使用Markdown提取器...")
                content = await self._extract_md_basic(file_path)
            elif file_extension in [".xlsx", ".xls"]:
                logger.debug("📊 使用Excel提取器...")
                content = await self._extract_xlsx_basic(file_path)
            else:
                logger.warning(f"⚠️ 不支持的文件类型: {file_extension}")
                content = f"文件类型 {file_extension} 需要marker-pdf支持，请安装: pip install marker-pdf"

            logger.debug(f"✅ 内容提取完成，长度: {len(content)} 字符")

            # 为Markdown文件创建更详细的MarkdownContent对象
            if file_extension == ".md":
                logger.debug("🔍 创建详细的Markdown内容对象...")
                result = self._create_markdown_content_from_md(content, file_extension)
            else:
                logger.debug("📦 创建基础内容对象...")
                # 创建基础的MarkdownContent对象
                result = MarkdownContent(
                    text=content,
                    images={},
                    metadata={
                        "extraction_method": "basic",
                        "file_type": file_extension,
                    },
                    tables=[],
                    headers=[],
                    links=[],
                    code_blocks=[],
                    math_expressions=[],
                )

            extraction_time = time.time() - start_time
            logger.success(f"✅ 基础文本提取完成，耗时: {extraction_time:.2f}秒")
            return result

        except Exception as e:
            extraction_time = time.time() - start_time
            logger.error(
                f"❌ 基础文本提取失败，耗时: {extraction_time:.2f}秒，错误: {str(e)}"
            )
            raise Exception(f"基础文本提取失败: {str(e)}")

    async def _extract_pdf_basic(self, file_path: Path) -> str:
        """基础PDF文本提取"""
        try:
            import PyPDF2

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content.append(page.extract_text())
                return "\n".join(content)
        except ImportError:
            return "PDF处理需要PyPDF2或marker-pdf，请安装: pip install PyPDF2 或 pip install marker-pdf"
        except Exception as e:
            raise Exception(f"PDF提取失败: {str(e)}")

    async def _extract_docx_basic(self, file_path: Path) -> str:
        """基础DOCX文本提取"""
        try:
            from docx import Document

            doc = Document(file_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            return "\n".join(content)
        except ImportError:
            return "DOCX处理需要python-docx，请安装: pip install python-docx"
        except Exception as e:
            raise Exception(f"DOCX提取失败: {str(e)}")

    async def _extract_txt_basic(self, file_path: Path) -> str:
        """基础TXT文本提取"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    return f.read()
            except Exception as e:
                raise Exception(f"文本文件编码错误: {str(e)}")

    async def _extract_md_basic(self, file_path: Path) -> str:
        """基础Markdown文本提取"""
        logger.debug(f"📝 开始提取Markdown内容: {file_path.name}")

        try:
            # 首先尝试UTF-8编码读取
            logger.debug("🔤 尝试UTF-8编码读取...")
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            logger.debug(f"✅ UTF-8编码读取成功，内容长度: {len(content)} 字符")

            # 对Markdown内容进行基础处理，保持格式
            logger.debug("🔧 处理Markdown内容格式...")
            processed_content = self._process_markdown_content(content)
            logger.debug("✅ Markdown内容处理完成")

            return processed_content

        except UnicodeDecodeError:
            logger.warning("⚠️ UTF-8编码失败，尝试GBK编码...")
            try:
                # 如果UTF-8失败，尝试GBK编码
                with open(file_path, "r", encoding="gbk") as f:
                    content = f.read()

                logger.debug(f"✅ GBK编码读取成功，内容长度: {len(content)} 字符")

                processed_content = self._process_markdown_content(content)
                logger.debug("✅ Markdown内容处理完成")

                return processed_content
            except Exception as e:
                logger.error(f"❌ GBK编码也失败: {str(e)}")
                raise Exception(f"Markdown文件编码错误: {str(e)}")
        except Exception as e:
            logger.error(f"❌ Markdown文件读取失败: {str(e)}")
            raise Exception(f"Markdown文件读取失败: {str(e)}")

    def _process_markdown_content(self, content: str) -> str:
        """处理Markdown内容，保持格式并进行基础优化"""
        # 移除多余的空行，但保持Markdown结构
        lines = content.split("\n")
        processed_lines = []
        prev_empty = False

        for line in lines:
            stripped_line = line.strip()

            # 如果是空行
            if not stripped_line:
                # 避免连续多个空行
                if not prev_empty:
                    processed_lines.append("")
                prev_empty = True
            else:
                processed_lines.append(line)
                prev_empty = False

        # 确保文档以单个换行结束
        processed_content = "\n".join(processed_lines).strip()

        return processed_content

    def _create_markdown_content_from_md(
        self, content: str, file_extension: str
    ) -> MarkdownContent:
        """为Markdown文件创建详细的MarkdownContent对象"""
        # 使用MarkdownExtractor来分析Markdown内容
        temp_markdown_content = MarkdownContent(
            text=content,
            images={},
            metadata={
                "extraction_method": "markdown_native",
                "file_type": file_extension,
            },
            tables=[],
            headers=[],
            links=[],
            code_blocks=[],
            math_expressions=[],
        )

        # 使用MarkdownExtractor分析内容
        analyzed_content = self.markdown_extractor._analyze_markdown_content(
            content,
            {},
            {"extraction_method": "markdown_native", "file_type": file_extension},
        )

        return analyzed_content

    async def _process_markdown_directly(self, file_path: Path) -> MarkdownContent:
        """直接处理Markdown文件，无需转换"""
        logger.info(f"📝 开始直接处理Markdown文件: {file_path.name}")

        try:
            # 读取文件内容
            logger.debug("📖 读取文件内容...")
            content = await self._extract_md_basic(file_path)
            logger.debug(f"✅ 文件内容读取完成，长度: {len(content)} 字符")

            # 创建MarkdownContent对象
            logger.debug("🔍 分析Markdown结构...")
            markdown_content = self._create_markdown_content_from_md(content, ".md")

            # 记录分析结果
            logger.info(f"📊 Markdown结构分析完成:")
            logger.info(f"  - 标题数量: {len(markdown_content.headers)}")
            logger.info(f"  - 表格数量: {len(markdown_content.tables)}")
            logger.info(f"  - 链接数量: {len(markdown_content.links)}")
            logger.info(f"  - 代码块数量: {len(markdown_content.code_blocks)}")
            logger.info(f"  - 数学表达式数量: {len(markdown_content.math_expressions)}")

            # 如果有图片分析器且图片不为空，进行图片分析
            if markdown_content.images and self.image_analyzer.is_available():
                logger.info("🖼️ 检测到图片，开始进行图片分析...")
                analyzed_content = await self._analyze_images_in_content(
                    markdown_content
                )
                return analyzed_content

            return markdown_content

        except Exception as e:
            logger.error(f"❌ Markdown文件直接处理失败: {str(e)}")
            raise

    async def _analyze_images_in_content(
        self, markdown_content: MarkdownContent
    ) -> MarkdownContent:
        """分析内容中的图片并添加描述"""
        logger.info(f"🔍 开始分析 {len(markdown_content.images)} 张图片")

        if not markdown_content.images:
            logger.debug("📷 没有图片需要分析")
            return markdown_content

        if not self.image_analyzer.is_available():
            logger.warning("⚠️ 图片分析器不可用，跳过图片分析")
            return markdown_content

        try:
            # 批量分析图片
            analysis_results = await self.image_analyzer.analyze_images_batch(
                markdown_content.images
            )

            # 提取成功的描述
            image_descriptions = {}
            for image_name, result in analysis_results.items():
                if result.get("success", False):
                    image_descriptions[image_name] = result.get("description", "")
                    logger.debug(f"   ✅ {image_name}: 分析成功")
                else:
                    logger.warning(
                        f"   ❌ {image_name}: 分析失败 - {result.get('error', '未知错误')}"
                    )

            # 如果有成功的分析结果，替换文本中的图片引用
            if image_descriptions:
                logger.info(f"📝 将 {len(image_descriptions)} 个图片描述集成到文本中")
                enhanced_text = self.image_analyzer.replace_images_with_descriptions(
                    markdown_content.text, image_descriptions
                )

                # 创建新的MarkdownContent对象，包含图片描述
                enhanced_content = MarkdownContent(
                    text=enhanced_text,
                    images=markdown_content.images,
                    metadata={
                        **markdown_content.metadata,
                        "image_analysis": {
                            "total_images": len(markdown_content.images),
                            "analyzed_images": len(image_descriptions),
                            "analysis_results": analysis_results,
                            "analyzer_model": self.image_analyzer.model,
                        },
                    },
                    tables=markdown_content.tables,
                    headers=markdown_content.headers,
                    links=markdown_content.links,
                    code_blocks=markdown_content.code_blocks,
                    math_expressions=markdown_content.math_expressions,
                )

                # 重新分析增强后的内容
                logger.debug("🔄 重新分析增强后的内容结构...")
                final_content = self.markdown_extractor._analyze_markdown_content(
                    enhanced_text, enhanced_content.images, enhanced_content.metadata
                )

                logger.success(
                    f"✅ 图片分析完成，成功分析 {len(image_descriptions)} 张图片"
                )
                return final_content
            else:
                logger.warning("⚠️ 没有成功分析的图片")
                return markdown_content

        except Exception as e:
            logger.error(f"❌ 图片分析过程失败: {e}")
            # 分析失败时返回原始内容
            return markdown_content

    async def _extract_xlsx_basic(self, file_path: Path) -> str:
        """基础XLSX文本提取"""
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(file_path)
            content = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"工作表: {sheet_name}")
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_data):
                        content.append("\t".join(row_data))
                content.append("")
            return "\n".join(content)
        except ImportError:
            return "Excel处理需要openpyxl，请安装: pip install openpyxl"
        except Exception as e:
            raise Exception(f"Excel提取失败: {str(e)}")

    async def process_upload_file(self, file: UploadFile) -> Dict[str, Any]:
        """处理上传的文件"""
        start_time = time.time()
        logger.info(f"📤 开始处理上传文件: {file.filename}")

        # 验证文件类型
        logger.debug("🔍 验证文件类型...")
        if not self.config_manager.is_file_supported(file.filename):
            supported_types = ", ".join(self.config_manager.get_supported_extensions())
            logger.error(f"❌ 不支持的文件类型: {Path(file.filename).suffix}")
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {Path(file.filename).suffix}。支持的类型: {supported_types}",
            )

        logger.debug(f"✅ 文件类型验证通过: {Path(file.filename).suffix}")

        # 验证文件大小
        logger.debug("📏 验证文件大小...")
        file_content = await file.read()
        file_size_mb = len(file_content) / 1024 / 1024
        max_size_mb = self.config_manager.config.max_file_size / 1024 / 1024

        logger.info(
            f"📊 上传文件大小: {file_size_mb:.2f}MB (限制: {max_size_mb:.1f}MB)"
        )

        if len(file_content) > self.config_manager.config.max_file_size:
            logger.error(
                f"❌ 文件大小超过限制: {file_size_mb:.2f}MB > {max_size_mb:.1f}MB"
            )
            raise HTTPException(
                status_code=400, detail=f"文件大小超过限制 ({max_size_mb:.1f}MB)"
            )

        logger.debug("✅ 文件大小验证通过")

        # 重置文件指针
        await file.seek(0)

        # 生成唯一文件名
        file_extension = Path(file.filename).suffix.lower()
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = self.upload_dir / unique_filename

        logger.debug(f"💾 生成唯一文件名: {unique_filename}")
        logger.debug(f"📁 保存路径: {file_path}")

        try:
            # 保存文件
            logger.debug("💾 保存上传文件到本地...")
            with open(file_path, "wb") as f:
                f.write(file_content)

            logger.debug(f"✅ 文件保存成功: {file_path}")

            # 处理文件
            logger.info("🔄 开始处理保存的文件...")
            markdown_content = await self.process_file(file_path)

            # 获取结构化内容
            logger.debug("📊 生成结构化内容...")
            structured_content = self.markdown_extractor.get_structured_content(
                markdown_content
            )

            processing_time = time.time() - start_time
            logger.success(f"✅ 上传文件处理完成，总耗时: {processing_time:.2f}秒")

            # 记录处理结果统计
            stats = structured_content["statistics"]
            logger.info(f"📈 处理结果统计: {structured_content}")
            logger.info(f"  - 字符数: {stats['total_characters']}")
            logger.info(f"  - 词数: {stats['total_words']}")
            logger.info(f"  - 标题: {stats['headers_count']} 个")
            logger.info(f"  - 表格: {stats['tables_count']} 个")
            logger.info(f"  - 链接: {stats['links_count']} 个")
            logger.info(f"  - 代码块: {stats['code_blocks_count']} 个")

            return {
                "status": "success",
                "filename": file.filename,
                "saved_filename": unique_filename,
                "file_path": str(file_path),
                "file_size": len(file_content),
                "file_type": file_extension,
                "markdown_content": structured_content,
                "processing_config": self.config_manager.get_config_dict(),
                "processing_time": processing_time,
            }

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(
                f"❌ 上传文件处理失败，耗时: {processing_time:.2f}秒，错误: {str(e)}"
            )

            # 如果处理失败，删除已保存的文件
            if file_path.exists():
                logger.debug("🗑️ 清理失败的临时文件...")
                file_path.unlink()
                logger.debug("✅ 临时文件清理完成")

            raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")

    async def process_multiple_files(
        self, file_paths: List[Union[str, Path]]
    ) -> List[Dict[str, Any]]:
        """批量处理多个文件"""
        results = []

        for file_path in file_paths:
            try:
                markdown_content = await self.process_file(file_path)
                structured_content = self.markdown_extractor.get_structured_content(
                    markdown_content
                )

                results.append(
                    {
                        "status": "success",
                        "file_path": str(file_path),
                        "markdown_content": structured_content,
                    }
                )

            except Exception as e:
                results.append(
                    {"status": "error", "file_path": str(file_path), "error": str(e)}
                )

        return results

    def update_config(self, **kwargs) -> None:
        """更新配置"""
        self.config_manager.update_config(**kwargs)
        # 重置converter以使用新配置
        self._converter = None

    def get_config(self) -> Dict[str, Any]:
        """获取当前配置"""
        return self.config_manager.get_config_dict()

    def search_in_content(
        self,
        markdown_content: MarkdownContent,
        query: str,
        case_sensitive: bool = False,
    ) -> List[Dict[str, Any]]:
        """在内容中搜索"""
        return self.markdown_extractor.search_content(
            markdown_content, query, case_sensitive
        )

    def save_content(self, markdown_content: MarkdownContent, output_path: str) -> None:
        """保存内容到文件"""
        self.markdown_extractor.save_content(markdown_content, output_path)

    def get_file_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """获取文件信息"""
        try:
            path = Path(file_path)
            if not path.exists():
                return None

            stat = path.stat()
            return {
                "filename": path.name,
                "file_size": stat.st_size,
                "file_extension": path.suffix.lower(),
                "is_supported": self.config_manager.is_file_supported(path.name),
                "created_time": stat.st_ctime,
                "modified_time": stat.st_mtime,
            }
        except Exception:
            return None

    def delete_file(self, file_path: str) -> bool:
        """删除文件"""
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                return True
            return False
        except Exception:
            return False

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """获取支持的文件格式信息"""
        extensions = self.config_manager.get_supported_extensions()

        format_groups = {
            "文档": [".pdf", ".docx", ".doc", ".html", ".htm", ".epub", ".md", ".txt"],
            "演示文稿": [".pptx", ".ppt"],
            "电子表格": [".xlsx", ".xls"],
            "图片": [".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"],
        }

        supported_formats = {}
        for group, group_extensions in format_groups.items():
            supported_formats[group] = [
                ext for ext in group_extensions if ext in extensions
            ]

        return supported_formats


class AsyncFileProcessor(FileProcessor):
    """异步文件处理器"""

    def __init__(
        self,
        config: Optional[MarkerConfig] = None,
        upload_dir: str = "uploads",
        max_concurrent: int = 3,
        image_analyzer: Optional[ImageAnalyzer] = None,
    ):
        super().__init__(config, upload_dir, image_analyzer)
        self.max_concurrent = max_concurrent
        self._semaphore = asyncio.Semaphore(max_concurrent)

    async def process_file_with_limit(
        self, file_path: Union[str, Path]
    ) -> MarkdownContent:
        """带并发限制的文件处理"""
        async with self._semaphore:
            return await self.process_file(file_path)

    async def process_multiple_files_concurrent(
        self, file_paths: List[Union[str, Path]]
    ) -> List[Dict[str, Any]]:
        """并发处理多个文件"""
        tasks = []

        for file_path in file_paths:
            task = asyncio.create_task(self._process_single_file_safe(file_path))
            tasks.append(task)

        results = await asyncio.gather(*tasks, return_exceptions=True)

        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append(
                    {
                        "status": "error",
                        "file_path": str(file_paths[i]),
                        "error": str(result),
                    }
                )
            else:
                processed_results.append(result)

        return processed_results

    async def _process_single_file_safe(
        self, file_path: Union[str, Path]
    ) -> Dict[str, Any]:
        """安全地处理单个文件"""
        try:
            markdown_content = await self.process_file_with_limit(file_path)
            structured_content = self.markdown_extractor.get_structured_content(
                markdown_content
            )

            return {
                "status": "success",
                "file_path": str(file_path),
                "markdown_content": structured_content,
            }
        except Exception as e:
            return {"status": "error", "file_path": str(file_path), "error": str(e)}


if __name__ == "__main__":
    file_process = AsyncFileProcessor(
        config=MarkerConfig(),
        upload_dir="/Users/bytedance/PycharmProjects/my_best/AITestLab/uploads",
    )
